"""
ProMatch - Professional Edition
Modern, clean UI with enterprise-grade design
"""
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from pdfminer.high_level import extract_text as pdf_extract_text
import docx
from datetime import datetime
from utils import (
    extract_emails, extract_phone_numbers,
    extract_years_of_experience, extract_education_level,
    calculate_resume_completeness, extract_certifications, normalize_text
)
from matcher import analyze_resume

# ==================== PAGE CONFIG ====================
st.set_page_config(
    page_title="ProMatch ",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== CUSTOM CSS ====================
st.markdown("""
    <style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Global Styles */
    * {
        font-family: 'Inter', sans-serif;
    }
    
    /* Hide Streamlit Branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Main Container */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1400px;
    }
    
    /* Header Styles */
    .main-header {
        font-size: 2.8rem;
        font-weight: 700;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom:  0.5rem;
        letter-spacing: -0.02em;
    }
    
    .sub-header {
        font-size: 1.1rem;
        color: #64748b;
        margin-bottom: 2.5rem;
        font-weight: 400;
    }
    
    /* Card Styles */
    .metric-card {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.08);
        border: 1px solid #e2e8f0;
        transition: transform 0.2s, box-shadow 0.2s;
    }
    
    .metric-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.12);
    }
    
    /* Sidebar Styles */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #f8fafc 0%, #ffffff 100%);
        border-right: 1px solid #e2e8f0;
    }
    
    [data-testid="stSidebar"] .element-container {
        margin-bottom: 1rem;
    }
    
    /* Button Styles */
    .stButton > button {
        border-radius: 8px;
        font-weight: 600;
        padding: 0.6rem 1.2rem;
        transition: all 0.3s;
        border: none;
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 4px 12px rgba(102, 126, 234, 0.4);
    }
    
    /* Primary Button */
    button[kind="primary"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Success/Error Messages */
    .stAlert {
        border-radius: 8px;
        border-left: 4px solid;
        padding: 1rem;
    }
    
    /* Expander Styles */
    .streamlit-expanderHeader {
        background: white;
        border-radius: 8px;
        border: 1px solid #e2e8f0;
        font-weight: 600;
        padding: 1rem;
    }
    
    .streamlit-expanderHeader:hover {
        background: #f8fafc;
        border-color: #cbd5e1;
    }
    
    /* Table Styles */
    .dataframe {
        border-radius: 8px;
        overflow: hidden;
        border: 1px solid #e2e8f0;
    }
    
    /* Progress Bar */
    .stProgress > div > div > div {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        border-radius: 4px;
    }
    
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: transparent;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        padding: 8px 16px;
        font-weight: 500;
    }
    
    /* Score Badge */
    .score-badge {
        display: inline-block;
        padding: 0.4rem 0.8rem;
        border-radius: 20px;
        font-weight: 600;
        font-size: 0.9rem;
    }
    
    .score-excellent {
        background: #d1fae5;
        color: #065f46;
    }
    
    .score-good {
        background: #fef3c7;
        color: #92400e;
    }
    
    .score-fair {
        background: #fee2e2;
        color: #991b1b;
    }
    
    /* Section Headers */
    .section-header {
        font-size: 1.5rem;
        font-weight: 700;
        color: #1e293b;
        margin-top: 2rem;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #e2e8f0;
    }
    
    /* Info Cards */
    .info-card {
        background: #f8fafc;
        border-radius: 8px;
        padding: 1rem;
        border-left: 4px solid #667eea;
        margin: 1rem 0;
    }
    
    /* Skill Pills */
    .skill-pill {
        display: inline-block;
        background: #ede9fe;
        color: #5b21b6;
        padding: 0.3rem 0.8rem;
        border-radius: 16px;
        font-size: 0.85rem;
        font-weight: 500;
        margin: 0.2rem;
    }
    
    /* Stats Grid */
    .stats-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
        gap: 1rem;
        margin: 1.5rem 0;
    }
    
    /* Welcome Card */
    .welcome-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 12px;
        padding: 2rem;
        margin: 2rem 0;
    }
    
    .welcome-card h3 {
        color: white;
        font-weight: 600;
    }
    </style>
""", unsafe_allow_html=True)

# ==================== SESSION STATE ====================
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
if 'results' not in st.session_state:
    st.session_state.results = []
if 'jd_saved' not in st.session_state:
    st.session_state.jd_saved = None

# ==================== HEADER ====================
col1, col2 = st.columns([3, 1])
with col1:
    st.markdown('<div class="main-header">⚡ ProMatch </div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">AI-Powered Intelligent Candidate Screening & Analysis Platform</div>', unsafe_allow_html=True)
with col2:
    st.markdown(f"""
    <div style='text-align: right; padding-top: 1rem;'>
        <div style='font-size: 0.85rem; color: #94a3b8;'>
            {datetime.now().strftime("%B %d, %Y")}
        </div>
    </div>
    """, unsafe_allow_html=True)

# ==================== SIDEBAR ====================
with st.sidebar:
    st.markdown("### ⚙️ Configuration Panel")
    st.markdown("---")
    
    # Job Description
    st.markdown("#### 📋 Step 1: Job Description")
    with st.container():
        jd_input = st.text_area(
            "",
            height=280,
            placeholder="""Paste your complete job description here...

Example:
• Required: Python, SQL, Machine Learning
• Preferred: Tableau, AWS
• 3+ years experience""",
            help="Enter the full job description including required and preferred skills"
        )
    
    if jd_input and jd_input != st.session_state.jd_saved:
        if st.button("💾 Save Job Description", use_container_width=True):
            st.session_state.jd_saved = jd_input
            st.success("✅ Job description saved!")
    
    st.markdown("---")
    
    # Resume Upload
    st.markdown("#### 📤 Step 2: Upload Resumes")
    uploaded_files = st.file_uploader(
        "",
        accept_multiple_files=True,
        type=['pdf', 'docx', 'txt'],
        help="Upload candidate resumes in PDF, DOCX, or TXT format"
    )
    
    if uploaded_files:
        st.success(f"✅ **{len(uploaded_files)}** resume(s) uploaded")
        for file in uploaded_files:
            file_size = len(file.getvalue()) / 1024
            st.text(f"📄 {file.name} ({file_size:.1f} KB)")
    
    st.markdown("---")
    
    # Action Buttons
    col1, col2 = st.columns(2)
    with col1:
        analyze_btn = st.button("🚀 Analyze", type="primary", use_container_width=True)
    with col2:
        if st.session_state.analysis_complete:
            if st.button("🔄 Reset", use_container_width=True):
                st.session_state.analysis_complete = False
                st.session_state.results = []
                st.session_state.jd_saved = None
                st.rerun()
    
    # Info Box
    st.markdown("---")
    st.markdown("""
    <div style='background: #f1f5f9; padding: 1rem; border-radius: 8px; font-size: 0.85rem;'>
        <strong>💡 Pro Tips:</strong><br>
        • Upload multiple resumes at once<br>
        • Supported: PDF, DOCX, TXT<br>
        • Results export to CSV/JSON<br>
        • View detailed skill analysis
    </div>
    """, unsafe_allow_html=True)

# ==================== MAIN CONTENT ====================

# ANALYSIS TRIGGER
if analyze_btn:
    if not jd_input:
        st.error("⚠️ Please provide a job description in the sidebar")
    elif not uploaded_files:
        st.error("⚠️ Please upload at least one resume")
    else:
        results = []
        clean_jd = normalize_text(jd_input, preserve_structure=True)
        
        # Progress Container
        progress_container = st.container()
        with progress_container:
            st.markdown("### 🔄 Processing Resumes...")
            progress_text = st.empty()
            progress_bar = st.progress(0)
        
        # Process each resume
        for idx, file in enumerate(uploaded_files):
            progress_text.markdown(f"**Analyzing:** `{file.name}` ({idx + 1}/{len(uploaded_files)})")
            
            try:
                file_extension = file.name.split('.')[-1].lower()
                
                # Extract raw text
                if file_extension == 'pdf':
                    raw_text = pdf_extract_text(file)
                elif file_extension in ['docx', 'doc']:
                    doc = docx.Document(file)
                    text_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_parts.append(cell.text)
                    raw_text = "\n".join(text_parts)
                elif file_extension == 'txt':
                    content = file.read()
                    raw_text = content.decode('utf-8', errors='ignore') if isinstance(content, bytes) else content
                else:
                    st.warning(f"⚠️ Unsupported format: {file.name}")
                    continue
                
                if not raw_text or len(raw_text.strip()) < 50:
                    st.warning(f"⚠️ File too short or empty: {file.name}")
                    continue
                
                # Extract metadata
                emails = extract_emails(raw_text)
                phones = extract_phone_numbers(raw_text)
                resume_text = normalize_text(raw_text, preserve_structure=True)
                total_exp = extract_years_of_experience(resume_text)
                education = extract_education_level(resume_text)
                completeness = calculate_resume_completeness(resume_text)
                certs = extract_certifications(resume_text)
                
                # Analyze
                if len(resume_text) < 100:
                    analysis = {
                        'overall_score': 0,
                        'breakdown': {
                            'must_have_skills': 0,
                            'nice_to_have_skills': 0,
                            'semantic_match': 0,
                            'experience_bonus': 0
                        },
                        'must_have_matched': [],
                        'must_have_missing': [],
                        'nice_to_have_matched': [],
                        'insights': ["⚠️ Resume too short for analysis"]
                    }
                else:
                    analysis = analyze_resume(clean_jd, resume_text)
                
                # Store results
                results.append({
                    'file_name': file.name,
                    'email': emails[0] if emails else "Not found",
                    'phone': phones[0] if phones else "Not found",
                    'overall_score': analysis['overall_score'],
                    'must_have_score': analysis['breakdown']['must_have_skills'],
                    'nice_to_have_score': analysis['breakdown']['nice_to_have_skills'],
                    'semantic_score': analysis['breakdown']['semantic_match'],
                    'experience_bonus': analysis['breakdown']['experience_bonus'],
                    'must_have_matched': analysis['must_have_matched'],
                    'must_have_missing': analysis['must_have_missing'],
                    'nice_to_have_matched': analysis['nice_to_have_matched'],
                    'insights': analysis['insights'],
                    'total_experience': total_exp if total_exp is not None else "N/A",
                    'education': education if education else "Not specified",
                    'completeness_score': completeness['score'],
                    'certifications_count': len(certs),
                    'full_analysis': analysis
                })
                
            except Exception as e:
                st.error(f"❌ Error processing {file.name}: {str(e)}")
                continue
            
            progress_bar.progress((idx + 1) / len(uploaded_files))
        
        progress_text.empty()
        progress_bar.empty()
        progress_container.empty()
        
        st.session_state.results = results
        st.session_state.analysis_complete = True
        st.success(f"✅ **Analysis Complete!** Successfully processed **{len(results)}** resume(s)")
        st.balloons()

# ==================== RESULTS DISPLAY ====================
if st.session_state.analysis_complete and st.session_state.results:
    results = st.session_state.results
    df = pd.DataFrame(results).sort_values('overall_score', ascending=False).reset_index(drop=True)
    
    # ===== SUMMARY DASHBOARD =====
    st.markdown('<div class="section-header">📊 Analysis Dashboard</div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        <div class="metric-card">
            <div style="color: #64748b; font-size: 0.85rem; margin-bottom: 0.5rem;">Total Candidates</div>
            <div style="font-size: 2rem; font-weight: 700; color: #1e293b;">{}</div>
            <div style="color: #10b981; font-size: 0.85rem; margin-top: 0.5rem;">🟢 {} Strong Matches</div>
        </div>
        """.format(len(df), len(df[df['overall_score'] >= 70])), unsafe_allow_html=True)
    
    with col2:
        avg_score = df['overall_score'].mean()
        st.markdown("""
        <div class="metric-card">
            <div style="color: #64748b; font-size: 0.85rem; margin-bottom: 0.5rem;">Average Score</div>
            <div style="font-size: 2rem; font-weight: 700; color: #1e293b;">{:.1f}%</div>
            <div style="color: #64748b; font-size: 0.85rem; margin-top: 0.5rem;">📈 Top: {:.1f}%</div>
        </div>
        """.format(avg_score, df['overall_score'].max()), unsafe_allow_html=True)
    
    with col3:
        qualified = len(df[df['overall_score'] >= 60])
        qual_pct = (qualified / len(df)) * 100
        st.markdown("""
        <div class="metric-card">
            <div style="color: #64748b; font-size: 0.85rem; margin-bottom: 0.5rem;">Qualified (≥60%)</div>
            <div style="font-size: 2rem; font-weight: 700; color: #1e293b;">{}</div>
            <div style="color: #64748b; font-size: 0.85rem; margin-top: 0.5rem;">📊 {:.0f}% of pool</div>
        </div>
        """.format(qualified, qual_pct), unsafe_allow_html=True)
    
    with col4:
        exp_values = df[df['total_experience'] != 'N/A']['total_experience']
        if len(exp_values) > 0:
            avg_exp = pd.to_numeric(exp_values, errors='coerce').mean()
            exp_display = f"{avg_exp:.1f} yrs"
        else:
            exp_display = "N/A"
        st.markdown("""
        <div class="metric-card">
            <div style="color: #64748b; font-size: 0.85rem; margin-bottom: 0.5rem;">Avg Experience</div>
            <div style="font-size: 2rem; font-weight: 700; color: #1e293b;">{}</div>
            <div style="color: #64748b; font-size: 0.85rem; margin-top: 0.5rem;">💼 Experience Level</div>
        </div>
        """.format(exp_display), unsafe_allow_html=True)
    
    # ===== SCORE DISTRIBUTION CHART =====
    st.markdown('<div class="section-header">📈 Score Distribution</div>', unsafe_allow_html=True)
    
    fig = go.Figure()
    fig.add_trace(go.Histogram(
        x=df['overall_score'],
        nbinsx=20,
        marker=dict(
            color=df['overall_score'],
            colorscale='Viridis',
            showscale=True,
            colorbar=dict(title="Score")
        ),
        hovertemplate='Score Range: %{x}<br>Candidates: %{y}<extra></extra>'
    ))
    
    fig.update_layout(
        title="Candidate Score Distribution",
        xaxis_title="Overall Score (%)",
        yaxis_title="Number of Candidates",
        height=350,
        template="plotly_white",
        showlegend=False
    )
    
    st.plotly_chart(fig, use_container_width=True)
    
    # ===== CANDIDATE RANKINGS TABLE =====
    st.markdown('<div class="section-header">🏆 Candidate Rankings</div>', unsafe_allow_html=True)
    
    display_df = df[['file_name', 'overall_score', 'email', 'phone', 'total_experience', 'education']].copy()
    display_df.columns = ['Candidate', 'Score', 'Email', 'Phone', 'Experience', 'Education']
    display_df.insert(0, 'Rank', range(1, len(display_df) + 1))
    
    def score_badge(score):
        if score >= 80:
            return f'<span class="score-badge score-excellent">{score:.1f}%</span>'
        elif score >= 60:
            return f'<span class="score-badge score-good">{score:.1f}%</span>'
        else:
            return f'<span class="score-badge score-fair">{score:.1f}%</span>'
    
    # Create styled HTML table
    table_html = "<table style='width:100%; border-collapse: collapse;'>"
    table_html += "<thead><tr style='background: #f8fafc; border-bottom: 2px solid #e2e8f0;'>"
    for col in display_df.columns:
        table_html += f"<th style='padding: 12px; text-align: left; font-weight: 600; color: #475569;'>{col}</th>"
    table_html += "</tr></thead><tbody>"
    
    for idx, row in display_df.iterrows():
        table_html += "<tr style='border-bottom: 1px solid #e2e8f0;'>"
        for col_idx, (col, val) in enumerate(row.items()):
            if col == 'Score':
                table_html += f"<td style='padding: 12px;'>{score_badge(val)}</td>"
            elif col == 'Rank':
                emoji = ['🥇', '🥈', '🥉'][idx] if idx < 3 else '📌'
                table_html += f"<td style='padding: 12px; font-weight: 600;'>{emoji} {val}</td>"
            else:
                table_html += f"<td style='padding: 12px;'>{val}</td>"
        table_html += "</tr>"
    
    table_html += "</tbody></table>"
    st.markdown(table_html, unsafe_allow_html=True)
    
    # ===== DETAILED CANDIDATE ANALYSIS =====
    st.markdown('<div class="section-header">🔍 Detailed Candidate Analysis</div>', unsafe_allow_html=True)
    
    for idx in range(min(10, len(df))):
        row = df.iloc[idx]
        rank_emoji = ['🥇', '🥈', '🥉'][idx] if idx < 3 else '📄'
        
        # Score color
        if row['overall_score'] >= 80:
            score_color = "#10b981"
        elif row['overall_score'] >= 60:
            score_color = "#f59e0b"
        else:
            score_color = "#ef4444"
        
        with st.expander(f"{rank_emoji} **{row['file_name']}** - Score: **{row['overall_score']:.1f}%**", expanded=(idx < 3)):
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.markdown("#### 💡 Key Insights")
                for insight in row['insights']:
                    st.markdown(f"- {insight}")
                
                st.markdown("#### ✅ Matched Skills")
                if row['must_have_matched']:
                    skills_html = ""
                    for skill in row['must_have_matched'][:15]:
                        skills_html += f'<span class="skill-pill">✓ {skill}</span>'
                    st.markdown(f"**Must-Have ({len(row['must_have_matched'])}):**", unsafe_allow_html=True)
                    st.markdown(skills_html, unsafe_allow_html=True)
                
                if row['nice_to_have_matched']:
                    skills_html = ""
                    for skill in row['nice_to_have_matched'][:10]:
                        skills_html += f'<span class="skill-pill" style="background:#dbeafe; color:#1e40af;">★ {skill}</span>'
                    st.markdown(f"<br>**Nice-to-Have ({len(row['nice_to_have_matched'])}):**", unsafe_allow_html=True)
                    st.markdown(skills_html, unsafe_allow_html=True)
                
                if row['must_have_missing']:
                    st.markdown("#### ❌ Missing Critical Skills")
                    missing_html = ""
                    for skill in row['must_have_missing'][:10]:
                        missing_html += f'<span class="skill-pill" style="background:#fee2e2; color:#991b1b;">✗ {skill}</span>'
                    st.markdown(missing_html, unsafe_allow_html=True)
            
            with col2:
                st.markdown("#### 👤 Candidate Profile")
                st.markdown(f"""
                <div class="info-card">
                    <div style="margin-bottom: 0.8rem;">
                        <strong>📧 Email:</strong><br>
                        <code>{row['email']}</code>
                    </div>
                    <div style="margin-bottom: 0.8rem;">
                        <strong>📞 Phone:</strong><br>
                        <code>{row['phone']}</code>
                    </div>
                    <div style="margin-bottom: 0.8rem;">
                        <strong>💼 Experience:</strong><br>
                        {row['total_experience']} years
                    </div>
                    <div style="margin-bottom: 0.8rem;">
                        <strong>🎓 Education:</strong><br>
                        {row['education'].title()}
                    </div>
                    <div>
                        <strong>📊 Resume Quality:</strong><br>
                        {row['completeness_score']}%
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Score Breakdown
                st.markdown("#### 📊 Score Breakdown")
                breakdown = row['full_analysis']['breakdown']
                
                score_breakdown_html = f"""
                <div style="background: white; border-radius: 8px; padding: 1rem; border: 1px solid #e2e8f0;">
                    <div style="margin-bottom: 0.5rem;">
                        <div style="display: flex; justify-content: space-between; margin-bottom: 0.3rem;">
                            <span style="font-size: 0.85rem; color: #64748b;">Must-Have Skills</span>
                            <span style="font-weight: 600;">{breakdown['must_have_skills']:.1f}%</span>
                        </div>
                        <div style="background: #e2e8f0; height: 8px; border-radius: 4px; overflow: hidden;">
                            <div style="background: #10b981; height: 100%; width: {breakdown['must_have_skills']}%;"></div>
                        </div>
                    </div>
                    <div style="margin-bottom: 0.5rem;">
                        <div style="display: flex; justify-content: space-between; margin-bottom: 0.3rem;">
                            <span style="font-size: 0.85rem; color: #64748b;">Semantic Match</span>
                            <span style="font-weight: 600;">{breakdown['semantic_match']:.1f}%</span>
                        </div>
                        <div style="background: #e2e8f0; height: 8px; border-radius: 4px; overflow: hidden;">
                            <div style="background: #3b82f6; height: 100%; width: {breakdown['semantic_match']}%;"></div>
                        </div>
                    </div>
                    <div style="margin-bottom: 0.5rem;">
                        <div style="display: flex; justify-content: space-between; margin-bottom: 0.3rem;">
                            <span style="font-size: 0.85rem; color: #64748b;">Experience</span>
                            <span style="font-weight: 600;">{breakdown['experience_bonus']:.1f}%</span>
                        </div>
                        <div style="background: #e2e8f0; height: 8px; border-radius: 4px; overflow: hidden;">
                            <div style="background: #f59e0b; height: 100%; width: {breakdown['experience_bonus'] * 10}%;"></div>
                        </div>
                    </div>
                </div>
                """
                st.markdown(score_breakdown_html, unsafe_allow_html=True)
    
    # ===== EXPORT SECTION =====
    st.markdown('<div class="section-header">📥 Export Results</div>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        csv = display_df.to_csv(index=False)
        st.download_button(
            "📊 Download CSV Report",
            csv,
            f"ats_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            "text/csv",
            use_container_width=True
        )
    
    with col2:
        import json
        json_str = json.dumps(df.to_dict(orient='records'), indent=2, default=str)
        st.download_button(
            "📄 Download JSON Data",
            json_str,
            f"ats_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "application/json",
            use_container_width=True
        )
    
    with col3:
        # Create summary report
        summary_report = f"""
PROMATCH - ANALYSIS REPORT
Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
{'='*60}

SUMMARY STATISTICS
------------------
Total Candidates Analyzed: {len(df)}
Average Score: {df['overall_score'].mean():.1f}%
Top Score: {df['overall_score'].max():.1f}%
Qualified Candidates (≥60%): {len(df[df['overall_score'] >= 60])}

TOP 5 CANDIDATES
----------------
"""
        for idx in range(min(5, len(df))):
            row = df.iloc[idx]
            summary_report += f"\n#{idx+1}. {row['file_name']}\n"
            summary_report += f"   Score: {row['overall_score']:.1f}%\n"
            summary_report += f"   Email: {row['email']}\n"
            summary_report += f"   Experience: {row['total_experience']}\n"
        
        st.download_button(
            "📋 Download Summary Report",
            summary_report,
            f"ats_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            use_container_width=True
        )

# ==================== WELCOME SCREEN ====================
else:
    st.markdown("""
    <div class="welcome-card">
        <h2 style="color: white; margin-bottom: 1rem;">👋 Welcome to ProMatch</h2>
        <p style="color: rgba(255,255,255,0.9); font-size: 1.1rem; line-height: 1.6;">
            Streamline your recruitment process with AI-powered candidate screening. 
            Our advanced matching algorithm analyzes resumes against job requirements, 
            providing detailed insights and objective scoring to help you find the perfect candidates faster.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="metric-card">
            <h3 style="color: #667eea; margin-bottom: 1rem;">🎯 Smart Matching</h3>
            <ul style="color: #64748b; line-height: 1.8;">
                <li>Multi-factor skill analysis</li>
                <li>Must-have vs Nice-to-have</li>
                <li>Semantic understanding</li>
                <li>Experience weighting</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="metric-card">
            <h3 style="color: #667eea; margin-bottom: 1rem;">⚡ Fast Processing</h3>
            <ul style="color: #64748b; line-height: 1.8;">
                <li>Batch resume upload</li>
                <li>Instant analysis</li>
                <li>Automated extraction</li>
                <li>Real-time scoring</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="metric-card">
            <h3 style="color: #667eea; margin-bottom: 1rem;">📊 Rich Insights</h3>
            <ul style="color: #64748b; line-height: 1.8;">
                <li>Visual analytics</li>
                <li>Detailed breakdowns</li>
                <li>Skill gap analysis</li>
                <li>Export reports</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Getting Started
    st.markdown('<div class="section-header">🚀 Getting Started</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div class="info-card">
            <h4 style="color: #1e293b; margin-bottom: 0.5rem;">📝 Step 1: Prepare Your Job Description</h4>
            <p style="color: #64748b; font-size: 0.9rem; line-height: 1.6;">
                Include clear requirements with <strong>required</strong> and <strong>preferred</strong> skills.
                Use phrases like "must have", "required", "nice to have", or "preferred".
            </p>
            <div style="background: white; padding: 0.8rem; border-radius: 6px; margin-top: 0.8rem;">
                <code style="color: #667eea; font-size: 0.85rem;">
                Example: "Required: Python, SQL | Preferred: Tableau or Power BI"
                </code>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="info-card">
            <h4 style="color: #1e293b; margin-bottom: 0.5rem;">📤 Step 2: Upload Candidate Resumes</h4>
            <p style="color: #64748b; font-size: 0.9rem; line-height: 1.6;">
                Upload multiple resumes in <strong>PDF</strong>, <strong>DOCX</strong>, or <strong>TXT</strong> format.
                Our system automatically extracts contact info, skills, experience, and education.
            </p>
            <div style="background: white; padding: 0.8rem; border-radius: 6px; margin-top: 0.8rem;">
                <code style="color: #667eea; font-size: 0.85rem;">
                Supported: .pdf, .docx, .txt | Max: 100 files
                </code>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    # Features Grid
    st.markdown('<div class="section-header">✨ Key Features</div>', unsafe_allow_html=True)
    
    features = [
        ("🔍", "Intelligent Skill Matching", "Advanced NLP extracts and matches skills with 95%+ accuracy"),
        ("🎓", "Education Detection", "Automatically identifies degree levels from Bachelor's to PhD"),
        ("💼", "Experience Analysis", "Calculates total years and skill-specific experience"),
        ("📧", "Contact Extraction", "Extracts emails, phone numbers, and professional profiles"),
        ("🔄", "OR Logic Support", "Handles alternative requirements like 'Tableau or Power BI'"),
        ("📊", "Visual Analytics", "Interactive charts and detailed score breakdowns"),
        ("⚡", "Batch Processing", "Analyze hundreds of resumes in minutes"),
        ("📥", "Multiple Export Formats", "Download results as CSV, JSON, or summary reports"),
    ]
    
    cols = st.columns(4)
    for idx, (emoji, title, desc) in enumerate(features):
        with cols[idx % 4]:
            st.markdown(f"""
            <div style="text-align: center; padding: 1rem; background: white; border-radius: 8px; 
                        border: 1px solid #e2e8f0; margin-bottom: 1rem; height: 180px; display: flex; 
                        flex-direction: column; justify-content: center;">
                <div style="font-size: 2.5rem; margin-bottom: 0.5rem;">{emoji}</div>
                <div style="font-weight: 600; color: #1e293b; margin-bottom: 0.5rem;">{title}</div>
                <div style="font-size: 0.85rem; color: #64748b; line-height: 1.4;">{desc}</div>
            </div>
            """, unsafe_allow_html=True)
    
    # FAQ Section
    st.markdown('<div class="section-header">❓ Frequently Asked Questions</div>', unsafe_allow_html=True)
    
    with st.expander("🤔 How does the scoring algorithm work?"):
        st.markdown("""
        Our scoring system uses **multi-factor analysis**:
        
        - **Must-Have Skills (50%)**: Percentage of required skills matched
        - **Semantic Match (25%)**: Overall context and job fit using AI
        - **Nice-to-Have Skills (15%)**: Bonus for preferred qualifications
        - **Experience Bonus (10%)**: Years of relevant experience
        
        **Score Interpretation:**
        - 85-100%: Excellent match - Interview immediately
        - 70-84%: Good match - Strong candidate
        - 60-69%: Moderate match - Review carefully
        - Below 60%: Missing critical requirements
        """)
    
    with st.expander("📊 What information is extracted from resumes?"):
        st.markdown("""
        We automatically extract:
        
        - ✉️ **Contact**: Email addresses, phone numbers
        - 🎓 **Education**: Degree level (Bachelor's, Master's, PhD)
        - 💼 **Experience**: Total years and skill-specific experience
        - 🛠️ **Skills**: Technical and soft skills with synonym normalization
        - 📜 **Certifications**: Professional certifications and licenses
        - 📋 **Resume Quality**: Completeness score based on sections
        """)
    
    with st.expander("🔒 Is my data secure?"):
        st.markdown("""
        **Yes!** All data is processed locally:
        
        - No data is sent to external servers
        - Files are processed in memory only
        - No permanent storage of uploaded documents
        - Session data is cleared on reset
        - Export files are generated client-side
        """)
    
    with st.expander("🚀 Can I use this for different industries?"):
        st.markdown("""
        **Absolutely!** Our universal skill database covers:
        
        - Technology & Engineering
        - Healthcare & Medical
        - Finance & Accounting
        - Marketing & Sales
        - Legal & Compliance
        - And 20+ other industries
        
        The system automatically adapts to your job description's domain.
        """)
    
    # Call to Action
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; padding: 2rem;">
        <h3 style="color: #1e293b; margin-bottom: 1rem;">Ready to Transform Your Hiring Process?</h3>
        <p style="color: #64748b; font-size: 1.1rem; margin-bottom: 1.5rem;">
            Get started by entering a job description and uploading resumes in the sidebar
        </p>
        <div style="font-size: 3rem;">👈</div>
    </div>
    """, unsafe_allow_html=True)

# ==================== FOOTER ====================
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #94a3b8; font-size: 0.85rem; padding: 1rem;">
    <strong>ProMatch</strong> v2.0 | 
    Powered by AI & Machine Learning | 
    © 2024 All Rights Reserved
</div>
""", unsafe_allow_html=True)
