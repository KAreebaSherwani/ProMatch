"""
IMPROVED UTILITIES
Better email, phone, education, and experience extraction
"""
from pdfminer.high_level import extract_text as pdf_extract_text
import docx
import re
from config import SKILL_SYNONYMS, STOP_WORDS

def extract_emails(text):
    """
    Extract email addresses with validation
    Works on RAW text before normalization
    """
    # Comprehensive email pattern
    pattern = r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
    
    emails = set()
    found = re.findall(pattern, text)
    
    # Filter and validate
    for email in found:
        email_lower = email.lower().strip()
        
        # Skip test/dummy emails
        if any(domain in email_lower for domain in ['test.com', 'domain.com', 'email.com', 'company.com']):
            continue
        
        # Validate structure
        if '@' in email_lower and '.' in email_lower.split('@')[1]:
            # Remove trailing punctuation
            email_clean = re.sub(r'[.,;:]+$', '', email_lower)
            emails.add(email_clean)
    
    return sorted(list(emails))


def extract_phone_numbers(text):
    """
    Extract phone numbers with multiple format support
    Works on RAW text before normalization
    """
    patterns = [
        # International: +1-234-567-8900
        r'\+\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        # US: (234) 567-8900
        r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}',
        # US: 234-567-8900
        r'\d{3}[-.\s]\d{3}[-.\s]\d{4}',
        # US: 234.567.8900
        r'\d{3}\.\d{3}\.\d{4}',
        # International: +92 300 1234567
        r'\+\d{1,3}\s\d{3}\s\d{7}',
        # Simple: 2345678900
        r'\b\d{10}\b',
    ]
    
    phones = set()
    
    for pattern in patterns:
        found = re.findall(pattern, text)
        for phone in found:
            # Clean up
            phone_clean = phone.strip()
            # Skip if looks like a date or ID
            if not re.match(r'^\d{4}[-/]\d{2}[-/]\d{2}', phone_clean):
                phones.add(phone_clean)
    
    return sorted(list(phones))


def normalize_text(text, preserve_structure=False):
    """
    Normalize text with synonym replacement
    Call this AFTER extracting emails/phones
    """
    if not text:
        return ""
    
    text = text.lower()
    
    # Apply skill synonyms (sorted by length for proper replacement)
    sorted_skills = sorted(SKILL_SYNONYMS.items(), key=lambda x: len(x[0]), reverse=True)
    
    for abbr, full in sorted_skills:
        # Use word boundaries to avoid partial matches
        pattern = r'\b' + re.escape(abbr) + r'\b'
        text = re.sub(pattern, full, text, flags=re.IGNORECASE)
    
    # Remove URLs
    text = re.sub(r'http\S+|www\.\S+', '', text)
    
    # Remove email addresses (already extracted)
    text = re.sub(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b', '', text)
    
    # Keep meaningful special characters
    text = re.sub(r'[^a-z0-9\s.+#/\-_\n]', ' ', text)
    
    # Normalize whitespace
    if preserve_structure:
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
    else:
        text = re.sub(r'\s+', ' ', text)
    
    return text.strip()


def extract_text_from_file(uploaded_file):
    """
    Extract text from uploaded file (PDF, DOCX, TXT)
    Returns (text, error_message)
    """
    try:
        file_name = uploaded_file.name if hasattr(uploaded_file, 'name') else str(uploaded_file)
        file_extension = file_name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            text = pdf_extract_text(uploaded_file)
            
        elif file_extension in ['docx', 'doc']:
            doc = docx.Document(uploaded_file)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            
        elif file_extension == 'txt':
            if hasattr(uploaded_file, 'read'):
                content = uploaded_file.read()
                text = content.decode('utf-8', errors='ignore') if isinstance(content, bytes) else content
            else:
                with open(uploaded_file, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
        else:
            return "", f"Unsupported file format: {file_extension}"
        
        # Validate
        if not text or len(text.strip()) < 50:
            return "", "File appears empty or too short"
        
        return text, None
        
    except Exception as e:
        return "", f"Error processing file: {str(e)}"


def extract_years_of_experience(resume_text):
    """
    Extract total years of experience
    Multiple pattern matching for accuracy
    """
    patterns = [
        # Direct statements
        r'(\d+)\+?\s*years?\s+(?:of\s+)?(?:total\s+)?(?:professional\s+)?experience',
        r'experience[:\s]+(\d+)\+?\s*years?',
        r'total\s+experience[:\s]+(\d+)\+?\s*years?',
        
        # Contextual
        r'(?:over|more than|around)\s+(\d+)\s+years?',
        r'(\d+)\+?\s*years?\s+(?:in|of|with)\s+(?:the\s+)?(?:field|industry)',
        r'for\s+(\d+)\+?\s*years?',
        r'(\d+)\s+years?\s+(?:of\s+)?(?:relevant\s+)?experience',
    ]
    
    years = []
    text_lower = resume_text.lower()
    
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            try:
                y = int(match)
                if 0 < y <= 50:  # Reasonable range
                    years.append(y)
            except:
                pass
    
    # Return maximum found (most comprehensive statement)
    return max(years) if years else None


def extract_education_level(resume_text):
    """
    STRICT education level detection
    Only returns if EXPLICITLY mentioned
    NO false positives
    """
    text_lower = resume_text.lower()
    
    # Try to isolate education section first
    education_section = ""
    
    edu_section_patterns = [
        r'education\s*:?(.*?)(?:experience|skills|projects|certifications|achievements|$)',
        r'academic\s+(?:background|qualifications)\s*:?(.*?)(?:experience|skills|$)',
        r'qualifications\s*:?(.*?)(?:experience|skills|$)',
    ]
    
    for pattern in edu_section_patterns:
        match = re.search(pattern, text_lower, re.DOTALL | re.IGNORECASE)
        if match:
            education_section = match.group(1)
            break
    
    # If no section found, use full text but with STRICT matching
    search_text = education_section if education_section else text_lower
    
    # STRICT patterns - must have complete degree phrases
    education_patterns = {
        'phd': [
            r'\bph\.?d\.?\s+(?:in|of|degree)',
            r'\bdoctorate\s+(?:in|of|degree)',
            r'\bdoctoral\s+degree',
            r'\bdoctor\s+of\s+philosophy',
        ],
        'masters': [
            r'\bmaster[\'s]?\s+(?:degree\s+)?(?:in|of)\s+[a-z\s]+',
            r'\bm\.?s\.?\s+(?:in|of)\s+[a-z\s]+',
            r'\bm\.?b\.?a\.?\b',
            r'\bm\.?a\.?\s+(?:in|of)\s+[a-z\s]+',
            r'\bmsc\s+(?:in|of)\s+[a-z\s]+',
            r'\bmaster\s+of\s+[a-z\s]+',
        ],
        'bachelors': [
            r'\bbachelor[\'s]?\s+(?:degree\s+)?(?:in|of)\s+[a-z\s]+',
            r'\bb\.?s\.?\s+(?:in|of)\s+[a-z\s]+',
            r'\bb\.?a\.?\s+(?:in|of)\s+[a-z\s]+',
            r'\bb\.?tech\.?\b',
            r'\bb\.?e\.?\s+(?:in|of)\s+[a-z\s]+',
            r'\bbachelor\s+of\s+[a-z\s]+',
        ],
    }
    
    # Check in priority order (PhD > Masters > Bachelors)
    for level in ['phd', 'masters', 'bachelors']:
        for pattern in education_patterns[level]:
            if re.search(pattern, search_text, re.IGNORECASE):
                return level
    
    return None


def calculate_resume_completeness(resume_text):
    """
    Score resume completeness based on key sections
    Returns dict with score and details
    """
    score = 0
    sections_found = []
    
    # Define sections and their point values
    section_checks = {
        'contact': {
            'patterns': [r'@', r'\+?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', r'linkedin', r'github'],
            'points': 15,
            'name': 'Contact Information'
        },
        'experience': {
            'patterns': [r'experience|employment|work history|professional background'],
            'points': 30,
            'name': 'Work Experience'
        },
        'education': {
            'patterns': [r'education|degree|university|college|academic'],
            'points': 20,
            'name': 'Education'
        },
        'skills': {
            'patterns': [r'skills|technologies|competencies|technical skills|expertise'],
            'points': 25,
            'name': 'Skills'
        },
        'summary': {
            'patterns': [r'summary|objective|about|profile|professional summary'],
            'points': 10,
            'name': 'Summary/Objective'
        }
    }
    
    text_lower = resume_text.lower()
    
    # Check each section
    for section_key, section_data in section_checks.items():
        found = False
        for pattern in section_data['patterns']:
            if re.search(pattern, text_lower):
                found = True
                break
        
        if found:
            score += section_data['points']
            sections_found.append(section_data['name'])
    
    # Calculate missing sections
    all_sections = [data['name'] for data in section_checks.values()]
    sections_missing = [s for s in all_sections if s not in sections_found]
    
    return {
        'score': min(score, 100),
        'sections_found': sections_found,
        'sections_missing': sections_missing
    }


def extract_certifications(resume_text):
    """
    Extract certifications and professional qualifications
    """
    cert_patterns = [
        r'([a-z\s]+)\s+certified',
        r'certified\s+([a-z\s]+)',
        r'([a-z\s]+)\s+certification',
        r'certification\s+in\s+([a-z\s]+)',
    ]
    
    # Common certification keywords
    cert_keywords = [
        'aws certified', 'azure certified', 'google certified',
        'pmp', 'prince2', 'csm', 'scrum master',
        'cpa', 'cfa', 'cma', 'cia',
        'cissp', 'ceh', 'comptia', 'ccna', 'ccnp',
        'six sigma', 'lean', 'itil',
        'oracle certified', 'microsoft certified', 'salesforce certified',
        'cima', 'acca', 'cpa', 'frm'
    ]
    
    certifications = set()
    text_lower = resume_text.lower()
    
    # Method 1: Pattern-based extraction
    for pattern in cert_patterns:
        matches = re.findall(pattern, text_lower)
        for match in matches:
            cert = match.strip()
            if 3 < len(cert) < 60 and cert not in STOP_WORDS:
                certifications.add(cert)
    
    # Method 2: Keyword matching
    for keyword in cert_keywords:
        if keyword in text_lower:
            # Extract surrounding context
            pattern = rf'.{{0,30}}{re.escape(keyword)}.{{0,30}}'
            matches = re.findall(pattern, text_lower)
            for match in matches:
                cert = match.strip()
                if cert:
                    certifications.add(cert)
    
    # Clean up and deduplicate
    unique_certs = []
    seen = set()
    
    for cert in certifications:
        # Normalize
        cert_clean = re.sub(r'\s+', ' ', cert).strip()
        
        if cert_clean not in seen and len(cert_clean) > 3:
            seen.add(cert_clean)
            unique_certs.append(cert_clean)
    
    return unique_certs[:15]  # Top 15